"""
Utility functions for candidate management.
Includes CV text extraction from PDF and DOCX files.
"""

import os
import re
from typing import Optional


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = []
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            
            return '\n'.join(text)
    except ImportError:
        print("PyPDF2 not installed. Trying pdfplumber...")
        try:
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                text = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                
                return '\n'.join(text)
        except ImportError:
            print("pdfplumber not installed. PDF parsing unavailable.")
            return None
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text from DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return '\n'.join(text)
    except ImportError:
        print("python-docx not installed. DOCX parsing unavailable.")
        return None
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None


def extract_text_from_doc(file_path: str) -> Optional[str]:
    """
    Extract text from DOC file (legacy Word format).
    
    Args:
        file_path: Path to the DOC file
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        import textract
        
        text = textract.process(file_path).decode('utf-8')
        return text
    except ImportError:
        print("textract not installed. DOC parsing unavailable.")
        return None
    except Exception as e:
        print(f"Error extracting text from DOC: {e}")
        return None


def extract_cv_text(file_path: str) -> Optional[str]:
    """
    Extract text from CV file based on file extension.
    Supports PDF, DOCX, DOC, and TXT files.
    
    Args:
        file_path: Path to the CV file
        
    Returns:
        Extracted text or None if extraction fails
    """
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return None
    
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.doc':
        return extract_text_from_doc(file_path)
    elif ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading text file: {e}")
            return None
    else:
        print(f"Unsupported file format: {ext}")
        return None


def clean_cv_text(text: str) -> str:
    """
    Clean and normalize CV text.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    
    # Remove control characters
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    return text.strip()
